from pathlib import Path
from docx import Document

SRC = Path('metodos_cualitativos_marino_costeros_actualizado.md')
OUT = Path('metodos_cualitativos_marino_costeros.docx')


def md_to_docx(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding='utf-8').splitlines()
    doc = Document()

    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(line)

    doc.save(out_path)


if __name__ == '__main__':
    md_to_docx(SRC, OUT)
    print(f'DOCX generado: {OUT}')
